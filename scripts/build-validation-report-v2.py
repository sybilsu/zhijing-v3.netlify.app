#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build-validation-report-v2.py — 影像語料重建後（2026-08-09）之附錄 B 驗證結果報告 v2 (.docx)
版式比照「設計過程說明書」：Arial、1A1A1A / 7A7A7A 單色系、深色表頭。

用法：python3 scripts/build-validation-report-v2.py
輸出：validation/植徑v3_附錄B_S_attr_S_llm_驗證結果報告_v2.docx
"""
import json, os, math, statistics as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
VAL = os.path.join(ROOT, 'validation')
PHOTOS = os.path.join(ROOT, 'public', 'photos')
INK, GREY, WHITE = RGBColor(0x1A, 0x1A, 0x1A), RGBColor(0x7A, 0x7A, 0x7A), RGBColor(0xFF, 0xFF, 0xFF)
FONT = 'Arial'

P2 = json.load(open(os.path.join(VAL, 'phase2-results.json'), encoding='utf8'))
P3 = json.load(open(os.path.join(VAL, 'phase3-results.json'), encoding='utf8'))

# ---------------------------------------------------------------- helpers
def set_font(run, size=11, color=INK, bold=False, spacing=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if spacing:
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing))
        run._element.rPr.append(sp)

def para(doc, text='', size=11, color=INK, bold=False, before=0, after=6,
         align=None, spacing=None, line=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align: p.alignment = align
    if text:
        set_font(p.add_run(text), size, color, bold, spacing)
    return p

def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(text), 15, INK, True)
    bar = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:space'), '6'); b.set(qn('w:color'), '1A1A1A')
    bar.append(b); p._p.get_or_add_pPr().append(bar)
    return p

def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), 11.5, INK, True)
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.4
    set_font(p.add_run(text), size, INK)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def table(doc, headers, rows, widths=None, size=9.5, head_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, '1A1A1A')
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after = Pt(3)
        set_font(c.paragraphs[0].add_run(htxt), head_size, WHITE, True)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].paragraphs[0].paragraph_format.space_before = Pt(2)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.25
            set_font(cells[i].paragraphs[0].add_run(str(v)), size, INK)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def caption(doc, text):
    para(doc, text, size=8.5, color=GREY, after=14, line=1.3)

def kv_block(doc, pairs):
    t = doc.add_table(rows=0, cols=2)
    for k, v in pairs:
        cells = t.add_row().cells
        cells[0].width = Inches(1.55); cells[1].width = Inches(4.95)
        for c in cells: c.paragraphs[0].paragraph_format.space_after = Pt(2)
        set_font(cells[0].paragraphs[0].add_run(k), 9.5, GREY)
        set_font(cells[1].paragraphs[0].add_run(v), 11, INK, True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------------------------------------------------------------- 統計
def _rank(v):
    idx = sorted(range(len(v)), key=lambda i: -v[i]); r = [0]*len(v); i = 0
    while i < len(idx):
        j = i
        while j+1 < len(idx) and v[idx[j+1]] == v[idx[i]]: j += 1
        for k in range(i, j+1): r[idx[k]] = (i+j)/2 + 1
        i = j+1
    return r

def rho(a, b):
    if len(a) < 3: return None
    rx, ry = _rank(a), _rank(b); mx, my = st.mean(rx), st.mean(ry)
    num = sum((x-mx)*(y-my) for x, y in zip(rx, ry))
    dx = sum((x-mx)**2 for x in rx); dy = sum((y-my)**2 for y in ry)
    return None if dx == 0 or dy == 0 else round(num/math.sqrt(dx*dy), 3)

pairs2 = [(r['input'], c) for r in P2['runs'] for c in r['candidates']]
S_llm = [c['s_llm'] for _, c in pairs2]
S_attr = [c['s_attr'] for _, c in pairs2]
pooled = rho(S_attr, S_llm)
NPAIR = len(pairs2)
NRUN = len(P2['runs'])

comp_rho = {}
for comp in ['foliage', 'height', 'spread', 'ornament']:
    vals = [v for v in (rho([c['breakdown'][comp] for c in r['candidates']],
                            [c['s_llm'] for c in r['candidates']]) for r in P2['runs']) if v is not None]
    comp_rho[comp] = (len(vals), round(st.mean(vals), 3) if vals else None)

sds = [c['s_llm_sd'] for r in P3['runs'] for c in r['candidates']]
stars = [c['stars'] for _, c in pairs2]

# ---------------------------------------------------------------- 圖 1 散布圖
def cjk_font():
    import matplotlib
    from matplotlib import font_manager
    for fam in ['Noto Sans CJK TC', 'Noto Sans CJK JP', 'Noto Serif CJK TC']:
        if any(fam in f.name for f in font_manager.fontManager.ttflist):
            matplotlib.rcParams['font.family'] = fam; break
    matplotlib.rcParams['axes.unicode_minus'] = False

def make_scatter(path):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    cjk_font()
    fig, ax = plt.subplots(figsize=(6.4, 3.8), dpi=200)
    marks = ['o', 's', '^', 'D', 'v', 'P', 'X', '*', 'h', '<']
    for i, r in enumerate(P2['runs']):
        ax.scatter([c['s_attr'] for c in r['candidates']], [c['s_llm'] for c in r['candidates']],
                   s=34, marker=marks[i % 10], facecolors='none', edgecolors='#1A1A1A', linewidths=1.0,
                   label=f"{r['input']['id']}（ρ={r['rho']}）")
    ax.axhline(60, color='#B0B0B0', lw=0.8, ls='--')
    ax.text(91, 61, 'S_llm = 60', fontsize=7, color='#7A7A7A', ha='right')
    ax.set_xlabel('S_attr（規則式屬性加權評分）', fontsize=9)
    ax.set_ylabel('S_llm（VLM 評審）', fontsize=9)
    ax.set_xlim(48, 93); ax.set_ylim(0, 100)
    ax.tick_params(labelsize=8)
    for s in ['top', 'right']: ax.spines[s].set_visible(False)
    for s in ['left', 'bottom']: ax.spines[s].set_color('#B0B0B0')
    ax.legend(fontsize=6.3, frameon=False, loc='upper left', ncol=2, handletextpad=0.4, columnspacing=0.8)
    fig.tight_layout(); fig.savefig(path, bbox_inches='tight'); plt.close(fig)

# ---------------------------------------------------------------- 圖 2 接觸表（重建後語料）
def judge_photo(name):
    for suf in ['_habit', '_close', '', '_leaf']:
        f = os.path.join(PHOTOS, f'{name}{suf}.jpg')
        if os.path.exists(f): return f
    return None

def make_contact(path, names):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.image as mpimg
    cjk_font()
    cols, rows_n = 5, 5
    fig, axes = plt.subplots(rows_n, cols, figsize=(10, 11.5), dpi=150)
    for i, ax in enumerate(axes.flat):
        ax.axis('off')
        if i < len(names):
            f = judge_photo(names[i])
            if f:
                img = mpimg.imread(f)
                ax.imshow(img)
            ax.set_title(f'{i+1}　{names[i]}', fontsize=9, color='#1A1A1A', pad=3)
    fig.tight_layout(h_pad=1.2, w_pad=0.6)
    fig.savefig(path, bbox_inches='tight'); plt.close(fig)

SCATTER = os.path.join(VAL, 'fig-scatter-v2.png')
make_scatter(SCATTER)
ALL_NAMES = sorted({r['input']['id'] for r in P2['runs']} |
                   {c['name'] for _, c in pairs2})
# 25 種完整名單（含未被抽入者）＝具照片者
NAMES25 = sorted(['田代氏石斑木', '絡石', '綬草', '青葙', '台灣糯米條', '台灣馬醉木', '密毛爵床', '山月桃',
                  '山素英', '截葉胡枝子', '桃金孃', '水社野牡丹', '田代氏黃芩', '紫苞舌蘭', '野牡丹', '鈴木草',
                  '阿里山油菊', '高士拂澤蘭', '台灣金絲桃', '杜虹花', '毛胡枝子', '燈稱花', '蠅翼草', '麥門冬', '黃荊'])
CONTACT = os.path.join(VAL, 'fig-contact-sheet-v2.jpg')
make_contact(CONTACT, NAMES25)

# ---------------------------------------------------------------- 文件
doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(1.0)
sec.top_margin = sec.bottom_margin = Inches(1.0)
normal = doc.styles['Normal']
normal.font.name = FONT; normal.font.size = Pt(11); normal.font.color.rgb = INK
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

# ---- 封面
para(doc, '', after=40)
para(doc, '第一屆景觀 AI 應用競賽　AI × LANDSCAPE 2026', 9.5, GREY, spacing=20, after=2)
para(doc, '第二組｜設計構想產生', 9.5, GREY, spacing=20, after=26)
para(doc, '植徑　NativePlant AI v3', 16, INK, True, after=6)
para(doc, '附錄 B｜S_attr / S_llm 驗證結果報告（第二輪）', 22, INK, True, after=6)
para(doc, 'APPENDIX B — SIMILARITY SCORING VALIDATION REPORT, ROUND 2', 9.5, GREY, spacing=30, after=26)
para(doc, '完成影像語料重建後之複驗：效標鏈上游的兩項污染已排除，'
          '本輪結果得以直接檢驗 S_attr 排序本身，並據以更新系統的真實瓶頸定位。', 11, INK, after=30, line=1.6)

kv_block(doc, [
    ('對應作品', 'zhijing-v3.netlify.app'),
    ('對應章節', '設計過程說明書 §5.2 發散到收斂與方案評估優化'),
    ('驗證版本', 'v.3.2（plants_enriched.json，508 筆；影像語料 2026-08-09 重建）'),
    ('執行日期', '2026 年 8 月 9 日'),
    ('前輪報告', '2026 年 8 月 4 日版（本文件為其 §6 後續行動 P0–P1 之執行結果）'),
    ('文件性質', '實驗性 prototype 驗證，非最終效能宣稱'),
])

doc.add_page_break()

# ---- 摘要
h1(doc, '摘要　EXECUTIVE SUMMARY')
para(doc, f"前輪（2026-08-04）驗證指出兩項必須先修復的前置缺陷：程式層的類別詞彙不一致，"
          f"以及影像語料的取景與正確性問題。本輪已完成 25 種送審影像的全面換補與稽核，"
          f"並依前輪建議將 Phase 2 擴大至 {NRUN} 組輸入植株（{NPAIR} 組配對），"
          f"另就 {P3['summary']['n_pairs']} 組配對各重複評分 3 次檢驗評審一致性。結論如下：", after=10)
table(doc,
      ['驗證指標', '門檻', '本輪結果', '判定'],
      [['Spearman ρ（S_attr vs S_llm）', 'ρ ≥ 0.6 最低／ρ ≥ 0.7 目標',
        f"組內平均 {P2['summary']['rho_mean']}；全樣本 pooled {pooled}", '未通過'],
       ['Overlap@5', '參考指標（無門檻）', f"平均 {P2['summary']['overlap5_mean']} / 5", '參考'],
       ['S_llm 一致性 SD', 'SD ≤ 15', f"平均 {round(st.mean(sds),2)}（最大 {max(sds)}）", '通過'],
       ['評審主體誤判率', '前輪 27.5%（11/40）', f"0%（0/{NPAIR}）", '已修復']],
      widths=[2.3, 1.7, 1.9, 0.75])
caption(doc, '表 1　驗證指標總覽。影像換補後評審不再誤判主體，效標鏈已乾淨；ρ 仍未達門檻。')

para(doc, '關鍵判讀：本輪與前輪的差別不在數字高低，而在證據力。前輪的 ρ ≈ 0 建立在受污染的效標上，'
          '不具備推翻或支持 S_attr 的能力；本輪影像語料已重建、類別詞彙已正規化、評審理由中主體誤判歸零，'
          '效標可信之後 ρ 仍接近零，即為一個可信的陰性結果——'
          '目前以屬性文字關鍵字比對推導的 S_attr 排序，無法預測視覺形態相似度的排序。'
          '瓶頸定位自「資料與影像是否可用」更新為「S_attr 的比對機制與樣本設計」。', after=10)

# ---- 一
h1(doc, '一、驗證目的與範圍　SCOPE')
para(doc, '植徑以規則式的 S_attr 屬性加權評分為候選植物排序。此分數透明、可回溯，但其排序是否貼近設計者'
          '實際的視覺判斷，無法由分數本身自證。本驗證引入獨立的視覺語言模型（VLM-as-Judge）作為外部效標 S_llm，'
          '以等級相關係數檢驗兩者的一致性。本輪為前輪報告 §6 後續行動之執行與複驗。', after=8)
h2(doc, '本輪涵蓋範圍（對應前輪 §6 優先序）')
bullet(doc, 'P0　影像語料重建：25 種送審影像逐種換補 _habit／_close／_leaf 槽位，剔除主體錯誤與含人工物影像（§3）')
bullet(doc, 'P0　類別正規化 normCategory()：維持於驗證腳本內生效；稽核黑名單 BAD_PHOTOS 因照片補正而清空')
bullet(doc, f"P1　重跑 Phase 2：輸入植株自 5 組擴至 {NRUN} 組（{NPAIR} 組配對），重新判定 ρ")
bullet(doc, f"Phase 3　S_llm 一致性檢驗：{P3['summary']['n_pairs']} 組配對 × 重複 3 次 ＝ 45 次評分")
h2(doc, '本輪未執行範圍')
bullet(doc, 'Phase 4 權重校準（NNLS）與 Phase 5 Role 維度 A/B：ρ 未達 0.6 門檻，依既定原則續行暫停（§6）')
bullet(doc, 'Phase 6 人類專家比對：尚未執行，本輪結果使其優先序提升（§7）')

# ---- 二
h1(doc, '二、方法　METHOD')
para(doc, '評分結構、評審設計與統計指標與前輪完全相同，僅樣本規模與影像語料不同，以維持前後輪可比較性。', after=8)
h2(doc, '2.1　S_attr 評分結構（總分 100）')
table(doc, ['分項', '權重', '計算方式'],
      [['類別 Category', '50（硬性門檻）', '輸入植株與資料庫植株類別須完全相同，否則排除'],
       ['葉形 Foliage', '30', '27 項葉形關鍵字詞庫 → Jaccard 相似度'],
       ['高度 Height', '10', '高度區間與資料庫 HEIGHT 欄位的數值重疊比例'],
       ['展幅 Spread', '5', '展幅區間與資料庫 SPREAD 欄位的數值重疊比例'],
       ['觀賞 Ornament', '5', '26 項觀賞特徵關鍵字（花期＋結構期）→ Jaccard 相似度']],
      widths=[1.35, 1.35, 3.85])
caption(doc, '表 2　S_attr 分項權重，與前端 src/matching.js 的 scorePair() 逐行一致。'
             '色彩刻意排除於本分數之外，由 Step 02 色彩篩選器獨立處理。')

h2(doc, '2.2　S_llm 評審設計')
table(doc, ['項目', '設定'],
      [['評審模型', 'claude-haiku-4-5-20251001（Anthropic Messages API）'],
       ['溫度 temperature', '1（保留取樣變異，使 Phase 3 的一致性 SD 具意義）'],
       ['影像處理', '本地照片統一縮至長邊 768 px、JPEG 品質 82'],
       ['評分尺度', '90–100 可直接互換／75–89 高度相似／60–74 中度／40–59 低度／0–39 不相似'],
       ['輸出解析', '正則擷取 JSON；解析失敗自動重試至多 3 次']],
      widths=[1.5, 5.0])
caption(doc, '表 3　S_llm 評審參數，與前輪相同；提示詞明確要求不得以花色或葉色作為主要依據。')

h2(doc, '2.3　樣本設計')
para(doc, '採 leave-one-out 設計：自資料庫抽取一筆作為「輸入植株」，以其屬性模擬 Step 02 影像辨識的輸出，'
          '再於資料庫其餘物種中排序。候選池限縮為本地備有照片的 25 種，'
          '輸入植株以類別分層、固定亂數種子（seed = 20260804）抽樣，確保可重現。', after=8)
table(doc, ['項目', '數值'],
      [['資料庫總數', '508 種（plants_enriched.json）'],
       ['具本地照片、可進入 S_llm 評審者', '25 種（4.9%）'],
       ['Phase 2 輸入植株', f"{NRUN} 組｜" + '、'.join(r['input']['id'] for r in P2['runs'])],
       ['Phase 2 配對數', f"{NPAIR} 組（每組輸入取 S_attr 前 8 名候選）"],
       ['Phase 3 配對數', f"{P3['summary']['n_pairs']} 組 × 重複 3 次 ＝ 45 次評分"],
       ['本輪新增 API 評審次數', '約 110 次｜費用約 US$0.35（重複配對由落地快取供應）']],
      widths=[2.35, 4.15])
caption(doc, '表 4　樣本規模。Phase 3 每組僅取前 3 名，抽樣結果與 Phase 2 不完全相同。')

h2(doc, '2.4　統計指標')
bullet(doc, 'Spearman ρ：組內計算（每組輸入植株的 8 筆候選各自排名），同分採平均秩；另計算全樣本 pooled ρ 供對照')
bullet(doc, 'Overlap@5：S_attr 前 5 名與 S_llm 前 5 名的交集筆數')
bullet(doc, 'S_llm 一致性 SD：同一配對重複 3 次評分的樣本標準差')

# ---- 三 影像語料重建
h1(doc, '三、影像語料重建　IMAGE CORPUS REBUILD')
para(doc, '本輪最主要的前置工作。依 PHOTO SPEC 的槽位規則（_habit 全株／群落、_close 枝條中間尺度、_leaf 葉形特寫；'
          '驗證腳本依 _habit → _close → 無後綴 → _leaf 順序取圖），對 25 種送審影像逐種目視稽核、換補與轉檔'
          '（長邊縮至 1568 px 以內、JPEG q85、sRGB）。前輪列入黑名單的四種（青葙＝鳥類照、綬草＝植物版畫、'
          '絡石＝蛾占畫面、田代氏石斑木＝住宅街景修剪灌木）全部換為野外實拍之合格照片，黑名單清空。', after=8)
h2(doc, '3.1　重建後取景稽核')
table(doc, ['取景類型', '張數', '占比', '對形態比對的可用性'],
      [['全株／群落尺度', '18', '72%', '可判讀株形、質地與量體'],
       ['枝條／植株局部', '6', '24%', '部分可判讀（該 6 種暫無合格全株照可用）'],
       ['沿用舊圖（全株、低解析）', '1', '4%', '紫苞舌蘭：換補來源屬別存疑（豆蘭屬），棄用而沿用原圖'],
       ['花／果／單葉特寫', '0', '0%', '—'],
       ['主體非目標植物', '0', '0%', '—']],
      widths=[1.9, 0.65, 0.65, 3.3])
caption(doc, '表 5　25 張送審影像重建後的取景稽核（前輪：全株 16%／特寫 60%／主體錯誤 8%）。'
             '_habit 檔名與內容的名實不符率由 76% 降為 0%。')
bullet(doc, '僅枝條尺度的 6 種：山月桃、山素英、水社野牡丹、杜虹花、毛胡枝子、台灣糯米條，列入後續補拍清單')
bullet(doc, '換補時同步剔除不合規來源：含量尺之照片（台灣糯米條）、含手部之照片（毛胡枝子）、'
            '含攝影浮水印之照片（山素英，僅降級為 _close 使用）')
bullet(doc, '紫苞舌蘭沿用原全株照（500 × 375 px，低於規格下限 800 px），為現存唯一未達像素規格者')
if os.path.exists(CONTACT):
    doc.add_picture(CONTACT, width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, '圖 1　重建後 25 張送審影像接觸表（編號依物種中文名排序；每種取驗證腳本實際讀用之 _habit 影像）。')

# ---- 四 結果
h1(doc, '四、結果　RESULTS')
h2(doc, '4.1　Phase 2 主結果')
rows = []
for r in P2['runs']:
    rows.append([r['input']['id'], r['input']['category'], r['n_pool'], len(r['candidates']),
                 r['rho'], f"{r['overlap5']} / 5"])
rows.append(['平均', '—', '—', P2['summary']['n_pairs'], P2['summary']['rho_mean'],
             f"{P2['summary']['overlap5_mean']} / 5"])
table(doc, ['輸入植株', '類別', '候選池', '送審', 'Spearman ρ', 'Overlap@5'], rows,
      widths=[1.5, 0.8, 0.8, 0.7, 1.35, 1.35])
caption(doc, f"表 6　各組輸入植株的排序一致性。{NRUN} 組中 {sum(1 for r in P2['runs'] if (r['rho'] or 0) < 0)} 組 ρ 為負、"
             f"{sum(1 for r in P2['runs'] if (r['rho'] or 0) > 0.5)} 組高於 0.5，平均值接近零。")
para(doc, f"全樣本 pooled ρ ＝ {pooled}，略高於零但遠低於門檻；顯示效標修復後，"
          f"S_attr 的排序與 S_llm 的排序之間仍不存在穩定的單調關係。"
          f"惟各組落差大（-0.591 至 0.648），與候選池組成高度相關：候選彼此形態同質性高的組（如細葉灌木群）"
          f"排序較可預測，形態異質的組則接近隨機。", after=10)
doc.add_picture(SCATTER, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption(doc, f"圖 2　S_attr 對 S_llm 散布圖（{NPAIR} 組配對）。點群仍集中於 S_llm 60 分以下，"
             '但已出現可解讀的高分配對（見 §4.2）。')

h2(doc, '4.2　S_llm 分數分布')
table(doc, ['統計量', '數值'],
      [['樣本數', f"{len(S_llm)} 組配對"],
       ['最小值 / 最大值', f"{min(S_llm)} / {max(S_llm)}"],
       ['平均數 / 中位數', f"{round(st.mean(S_llm),1)} / {st.median(S_llm)}"],
       ['標準差', f"{round(st.pstdev(S_llm),1)}"],
       ['≥ 60 分（中度相似以上）', f"{sum(1 for v in S_llm if v>=60)} / {len(S_llm)}（{round(100*sum(1 for v in S_llm if v>=60)/len(S_llm))}%）"],
       ['≤ 45 分（低度相似以下）', f"{sum(1 for v in S_llm if v<=45)} / {len(S_llm)}（{round(100*sum(1 for v in S_llm if v<=45)/len(S_llm))}%）"]],
      widths=[2.4, 4.1])
caption(doc, '表 7　S_llm 分數分布。地板效應仍在：評審主要使用尺度下半段。'
             '但與前輪不同，本輪的高分配對在植物學上可解讀。')
para(doc, '本輪 S_llm ≥ 60 的配對如：田代氏石斑木 × 桃金孃（78）、燈稱花 × 田代氏石斑木（72）、'
          '燈稱花 × 桃金孃（68）——皆為細葉、密枝、質地相近的常綠灌木組合，替換邏輯成立；'
          '低分配對的理由則描述雙方真實的形態差異（葉形、花序形式、株型量體）。'
          '對照前輪充斥「圖 B 為鳥類，非植物」「誤判為蕨類」之評審理由，'
          '本輪理由中主體誤判為 0 筆，顯示低分主要反映真實差異，而非影像缺陷。'
          '在候選池僅 8–10 種、同類別下形態天然分歧的條件下，多數配對確實不相似，地板效應有其真實成分。', after=10)

h2(doc, '4.3　S_attr 分項與 S_llm 的相關性')
label = {'foliage': '葉形 Foliage（權重 30）', 'height': '高度 Height（權重 10）',
         'spread': '展幅 Spread（權重 5）', 'ornament': '觀賞 Ornament（權重 5）'}
table(doc, ['S_attr 分項', '可計算組數', '組內平均 ρ'],
      [[label[k], f"{v[0]} / {NRUN}", f"{v[1]:+.3f}" if v[1] is not None else '—'] for k, v in comp_rho.items()],
      widths=[2.6, 1.5, 2.4])
caption(doc, f"表 8　各分項與 S_llm 的組內相關。權重最大的葉形分項（30 分）相關為負；"
             f"觀賞分項在 {NRUN} 組中僅 {comp_rho['ornament'][0]} 組可計算（其餘全組同分），與前輪相同地不具鑑別力。")

h2(doc, '4.4　Phase 3　S_llm 一致性')
rows3 = []
for r in P3['runs']:
    for c in r['candidates']:
        rows3.append([r['input']['id'], c['name'], c['s_attr'],
                      ' / '.join(str(v) for v in c['s_llm_runs']), c['s_llm'], c['s_llm_sd']])
table(doc, ['輸入植株', '候選植株', 'S_attr', '三次評分', '平均', 'SD'], rows3,
      widths=[1.25, 1.25, 0.75, 1.25, 0.7, 0.6], size=9)
caption(doc, f"表 9　重複評分一致性。平均 SD ＝ {round(st.mean(sds),2)}（最大 {max(sds)}），低於門檻 15，"
             f"其中 {sum(1 for v in sds if v==0)} 組配對三次評分完全相同。")
para(doc, 'S_llm 通過一致性檢驗。本輪 SD（3.43）略高於前輪（1.51），主因是換補後影像資訊量增加、'
          '評審在相鄰分數帶（如 32／35／42）間的取樣變異增大，仍遠低於門檻，屬健康範圍。'
          '與前輪相同的保留意見依然成立：高一致性只說明「評審可重複」，不能單獨推論「評審有鑑別力」。', after=10)

# ---- 五 換補成效
h1(doc, '五、影像換補成效對照　BEFORE / AFTER')
table(doc, ['項目', '前輪（08-04）', '本輪（08-09）'],
      [['送審影像：全株／群落尺度占比', '16%', '72%'],
       ['送審影像：花果特寫占比', '60%', '0%'],
       ['送審影像：主體非目標植物', '2 張（8%）', '0 張'],
       ['_habit 檔名名實不符率', '76%', '0%'],
       ['評審理由主體誤判', '11 / 40 組（27.5%）', f"0 / {NPAIR} 組（0%）"],
       ['組內平均 ρ', '-0.027（5 組）', f"{P2['summary']['rho_mean']}（{NRUN} 組）"],
       ['全樣本 pooled ρ', '0.087', f'{pooled}'],
       ['Overlap@5 平均', '3.4 / 5', f"{P2['summary']['overlap5_mean']} / 5"],
       ['S_llm 一致性 SD', '1.51', f'{round(st.mean(sds),2)}']],
      widths=[2.7, 1.9, 1.9])
caption(doc, '表 10　前後輪對照。影像品質指標全面改善、效標誤判歸零；相關性指標僅微幅上移，仍未達門檻。')
para(doc, '解讀：換補相片達成了它能達成的目標——讓 S_llm 這把「尺」本身變得可信。'
          '它沒有（也不應該）讓 ρ 自動變好；ρ 在乾淨效標下維持近零，正是本輪最有價值的訊號，'
          '它把問題明確指向 S_attr 的比對機制本身。', after=10)

# ---- 六 發現與診斷
h1(doc, '六、發現與診斷　FINDINGS')
h2(doc, '6.1　葉形分項失效：權重最大的維度與視覺無關')
para(doc, f"葉形 Foliage 佔 S_attr 軟性區間 30／50，是實際決定排序的最大力量，"
          f"但其與 S_llm 的組內平均 ρ ＝ {comp_rho['foliage'][1]:+.3f}。機制上，"
          '葉形分數來自資料庫文字欄位的 27 項關鍵字 Jaccard 比對——它比的是「描述詞彙的交集」，'
          '不是「看起來像不像」。兩種葉形描述詞彙高度重疊的物種（同用「橢圓、對生、革質」），'
          '在全株尺度下的質地與量體可以完全不同；反之，視覺上可互換的細葉灌木（如燈稱花與田代氏石斑木，'
          'S_llm＝72）在詞彙上未必高分。文字代理（text proxy）與視覺目標之間的落差，是 ρ 近零的主要結構性原因。', after=8)
h2(doc, '6.2　觀賞分項退化為常數')
para(doc, f"觀賞 Ornament 分項在 {NRUN} 組中 {NRUN - comp_rho['ornament'][0]} 組全體候選同分（多為 0 分），僅 "
          f"{comp_rho['ornament'][0]} 組可計算相關。花期／結構期關鍵字在同類別物種間普遍缺乏交集，"
          '該分項實質上不參與排序，與前輪發現一致、尚未修復。', after=8)
h2(doc, '6.3　小候選池與窄分數帶的統計限制')
para(doc, f"候選池僅 8–10 種時，單組 ρ 的抽樣變異極大（本輪橫跨 -0.591 至 0.648）。"
          f"且 S_attr 分數集中於 {min(S_attr)}–{max(S_attr)} 的窄帶（標準差 {round(st.pstdev(S_attr),1)}，"
          f"50 分為類別門檻固定基底）；★ 星等分布為 ★2＝{stars.count(2)}、★3＝{stars.count(3)}、"
          f"★4＝{stars.count(4)}、★5＝{stars.count(5)}，過半集中於 ★3。"
          '分數帶窄使同分與近同分頻繁出現，排序訊號進一步被稀釋。根本解法是擴大照片物種數，讓候選池變大。', after=8)
h2(doc, '6.4　S_llm 尺度使用偏窄，建議改採成對比較')
para(doc, '評審實際輸出集中於 28–45 帶，尺度上半段幾乎未使用。後續可改用錨定範例（rubric anchoring）'
          '或直接改為成對比較（給兩位候選、問哪個更像參考植株），將絕對評分轉為相對排序，'
          '可望同時緩解地板效應與分數帶壓縮的問題。', after=10)

# ---- 七 結論
h1(doc, '七、結論　CONCLUSION')
para(doc, '本輪驗證的判定為：Spearman ρ 未達 0.6 的最低可接受門檻；S_llm 一致性 SD 通過門檻；'
          '影像語料重建完成，評審主體誤判歸零。', after=8)
para(doc, '與前輪本質不同的是：前輪的 ρ 不具證據力，本輪的 ρ 是在乾淨效標上取得的可信陰性結果。'
          '它說明目前的 S_attr——以文字關鍵字交集近似視覺相似——在小候選池上無法重現視覺排序。'
          '據此，Phase 4 權重校準繼續暫停：在分項與效標普遍近零相關的條件下，'
          'NNLS 只會擬合噪音；正確的下一步是先修比對機制（葉形維度的視覺化、觀賞分項的詞庫重建），'
          '並以人類專家（Phase 6）錨定 S_llm 的效度，再回頭談權重。', after=8)
para(doc, '就 prototype 的目的而言，本輪完成了驗證方法論的一次完整閉環：發現缺陷 → 修復語料 → 複驗 → '
          '把瓶頸從「資料可不可用」推進到「機制對不對」。這條可重跑、可稽核的驗證路徑本身，'
          '是比任何單一分數都更重要的產出。', after=10)

# ---- 八 後續行動
h1(doc, '八、後續行動　NEXT ACTIONS')
table(doc, ['優先序', '行動', '預期效果'],
      [['P0', '葉形比對機制改造：以影像側特徵（VLM 對照片輸出的結構化葉形描述）取代或補強資料庫文字關鍵字',
        '直接處理 §6.1 的文字代理落差，為 ρ 提升建立機制基礎'],
       ['P0', '重建觀賞 Ornament 詞庫或將其權重併入葉形／質地維度',
        '消除退化為常數的分項'],
       ['P1', 'Phase 6：以 2–3 位景觀設計師對 12–15 組配對評分，與 S_llm 求一致性',
        '錨定效標效度；為 Phase 4 提供可信的校準目標'],
       ['P1', 'S_llm 評審改為成對比較或加入錨定範例',
        '緩解地板效應與尺度壓縮（§6.4）'],
       ['P1', '補拍 6 種僅枝條尺度影像（山月桃、山素英、水社野牡丹、杜虹花、毛胡枝子、台灣糯米條）'
              '與紫苞舌蘭合格全株照',
        '語料完備度 100%'],
       ['P2', '擴充照片物種數（現 25／508），使單組候選池 ≥ 15',
        '降低單組 ρ 的抽樣變異（§6.3）'],
       ['P2', '機制修復且 ρ ≥ 0.6 後，啟動 Phase 4 權重校準（NNLS，保留樣本）與 Phase 5 Role 維度 A/B',
        '在可信效標上完成權重最佳化']],
      widths=[0.65, 3.35, 2.5])
caption(doc, '表 11　後續行動優先序（更新版）。前輪 P0 兩項已於本輪完成並自清單移除。')

# ---- 九 可重現性
h1(doc, '九、可重現性　REPRODUCIBILITY')
para(doc, '本報告的所有數值皆由下列指令產生，固定亂數種子後可完整重現：', after=8)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
set_font(p.add_run('node scripts/validate-similarity.mjs --inputs 10 --topk 8 --repeat 1 \\\n'
                   '     --model claude-haiku-4-5-20251001 --seed 20260804'), 9.5, INK)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
set_font(p.add_run('node scripts/validate-similarity.mjs --inputs 5 --topk 3 --repeat 3 \\\n'
                   '     --model claude-haiku-4-5-20251001 --seed 20260804'), 9.5, INK)
table(doc, ['檔案', '內容'],
      [['scripts/validate-similarity.mjs', '驗證腳本（S_attr 評分、抽樣、S_llm 呼叫、統計；BAD_PHOTOS 已清空）'],
       ['scripts/build-validation-report-v2.py', '本報告的產生腳本'],
       ['validation/phase2-results.json', 'Phase 2 完整結果（80 組配對，含每筆評審理由）'],
       ['validation/phase3-results.json', 'Phase 3 重複評分結果'],
       ['validation/phase2-pairs.csv', '80 組配對的分項明細（可直接開啟檢視）'],
       ['validation/.judge-cache.json', '評審結果落地快取，重跑不重複計費'],
       ['validation/archive_20260804/', '前輪（08-04）全部結果檔之封存'],
       ['public/photos_backup_20260804/', '換補前影像語料之完整備份'],
       ['8月送件/換補相片/', '換補影像原始來源（依 01–25 物種編號歸檔）']],
      widths=[2.7, 3.8])
caption(doc, '表 12　產出檔案清單。')

# ---- 十 AI 揭露
h1(doc, '十、AI 工具揭露　AI DISCLOSURE')
table(doc, ['角色', '模型 / 版本', '用途'],
      [['視覺相似度評審', 'claude-haiku-4-5-20251001', 'S_llm 效標，對雙圖配對給出 0–100 分'],
       ['影像辨識（系統內）', 'GPT-4o Vision', '本驗證以資料庫屬性模擬其輸出，未實際呼叫'],
       ['規則式評分', '無 AI（確定性程式碼）', 'S_attr，與 src/matching.js 邏輯一致'],
       ['影像稽核與槽位分配', '人工逐張目視（2026-08-09）', '換補相片之取景分類與 _habit／_close／_leaf 指派']],
      widths=[1.7, 2.2, 2.6])
para(doc, '完整提示詞、模型版本、溫度設定與每筆評審的原始回覆，均保存於 validation/phase2-results.json 與'
          ' scripts/validate-similarity.mjs，可供查驗。', 9.5, GREY, after=6)

out = os.path.join(VAL, '植徑v3_附錄B_S_attr_S_llm_驗證結果報告_v2.docx')
doc.save(out)
print('saved:', out)
