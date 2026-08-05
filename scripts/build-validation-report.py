#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build-validation-report.py — 由 validation/*.json 產出附錄 B 驗證結果報告 (.docx)
版式比照「3_植徑v3_設計過程說明書.docx」：Arial、1A1A1A / 7A7A7A 單色系、深色表頭。

用法：python3 scripts/build-validation-report.py
輸出：validation/植徑v3_附錄B_S_attr_S_llm_驗證結果報告.docx
"""
import json, os, math, statistics as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
VAL = os.path.join(ROOT, 'validation')
INK, GREY, WHITE = RGBColor(0x1A, 0x1A, 0x1A), RGBColor(0x7A, 0x7A, 0x7A), RGBColor(0xFF, 0xFF, 0xFF)
FONT = 'Arial'

P2 = json.load(open(os.path.join(VAL, 'phase2-results.json'), encoding='utf8'))
P3 = json.load(open(os.path.join(VAL, 'phase3-results.json'), encoding='utf8'))

# ---------------------------------------------------------------- helpers
def set_font(run, size=11, color=INK, bold=False, spacing=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if spacing:
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing))
        run._element.rPr.append(sp)

def para(doc, text='', size=11, color=INK, bold=False, before=0, after=6,
         align=None, spacing=None, line=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align: p.alignment = align
    if text:
        set_font(p.add_run(text), size, color, bold, spacing)
    return p

def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(text), 15, INK, True)
    bar = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6'); b.set(qn('w:space'), '6'); b.set(qn('w:color'), '1A1A1A')
    bar.append(b); p._p.get_or_add_pPr().append(bar)
    return p

def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), 11.5, INK, True)
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.4
    set_font(p.add_run(text), size, INK)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def table(doc, headers, rows, widths=None, size=9.5, head_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, '1A1A1A')
        c.paragraphs[0].paragraph_format.space_before = Pt(3)
        c.paragraphs[0].paragraph_format.space_after = Pt(3)
        set_font(c.paragraphs[0].add_run(htxt), head_size, WHITE, True)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].paragraphs[0].paragraph_format.space_before = Pt(2)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.25
            set_font(cells[i].paragraphs[0].add_run(str(v)), size, INK)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def caption(doc, text):
    para(doc, text, size=8.5, color=GREY, after=14, line=1.3)

def kv_block(doc, pairs):
    t = doc.add_table(rows=0, cols=2)
    for k, v in pairs:
        cells = t.add_row().cells
        cells[0].width = Inches(1.55); cells[1].width = Inches(4.95)
        for c in cells: c.paragraphs[0].paragraph_format.space_after = Pt(2)
        set_font(cells[0].paragraphs[0].add_run(k), 9.5, GREY)
        set_font(cells[1].paragraphs[0].add_run(v), 11, INK, True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------------------------------------------------------------- 統計
def _rank(v):
    idx = sorted(range(len(v)), key=lambda i: -v[i]); r = [0]*len(v); i = 0
    while i < len(idx):
        j = i
        while j+1 < len(idx) and v[idx[j+1]] == v[idx[i]]: j += 1
        for k in range(i, j+1): r[idx[k]] = (i+j)/2 + 1
        i = j+1
    return r

def rho(a, b):
    if len(a) < 3: return None
    rx, ry = _rank(a), _rank(b); mx, my = st.mean(rx), st.mean(ry)
    num = sum((x-mx)*(y-my) for x, y in zip(rx, ry))
    dx = sum((x-mx)**2 for x in rx); dy = sum((y-my)**2 for y in ry)
    return None if dx == 0 or dy == 0 else round(num/math.sqrt(dx*dy), 3)

pairs2 = [(r['input'], c) for r in P2['runs'] for c in r['candidates']]
S_llm = [c['s_llm'] for _, c in pairs2]
S_attr = [c['s_attr'] for _, c in pairs2]
pooled = rho(S_attr, S_llm)

comp_rho = {}
for comp in ['foliage', 'height', 'spread', 'ornament']:
    vals = [v for v in (rho([c['breakdown'][comp] for c in r['candidates']],
                            [c['s_llm'] for c in r['candidates']]) for r in P2['runs']) if v is not None]
    comp_rho[comp] = (len(vals), round(st.mean(vals), 3) if vals else None)

sds = [c['s_llm_sd'] for r in P3['runs'] for c in r['candidates']]

# ---------------------------------------------------------------- 圖
def make_scatter(path):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for fam in ['Noto Sans CJK TC', 'Noto Sans CJK JP', 'Noto Serif CJK TC']:
        if any(fam in f.name for f in font_manager.fontManager.ttflist):
            matplotlib.rcParams['font.family'] = fam; break
    matplotlib.rcParams['axes.unicode_minus'] = False
    fig, ax = plt.subplots(figsize=(6.4, 3.5), dpi=200)
    marks = ['o', 's', '^', 'D', 'v']
    for i, r in enumerate(P2['runs']):
        ax.scatter([c['s_attr'] for c in r['candidates']], [c['s_llm'] for c in r['candidates']],
                   s=34, marker=marks[i % 5], facecolors='none', edgecolors='#1A1A1A', linewidths=1.0,
                   label=f"{r['input']['id']}（ρ={r['rho']}）")
    ax.axhline(60, color='#B0B0B0', lw=0.8, ls='--')
    ax.text(91, 61, 'S_llm = 60', fontsize=7, color='#7A7A7A', ha='right')
    ax.set_xlabel('S_attr（規則式屬性加權評分）', fontsize=9)
    ax.set_ylabel('S_llm（VLM 評審）', fontsize=9)
    ax.set_xlim(48, 93); ax.set_ylim(0, 100)
    ax.tick_params(labelsize=8)
    for s in ['top', 'right']: ax.spines[s].set_visible(False)
    for s in ['left', 'bottom']: ax.spines[s].set_color('#B0B0B0')
    ax.legend(fontsize=7, frameon=False, loc='upper left', ncol=2, handletextpad=0.4, columnspacing=0.9)
    fig.tight_layout(); fig.savefig(path, bbox_inches='tight'); plt.close(fig)

SCATTER = os.path.join(VAL, 'fig-scatter.png')
make_scatter(SCATTER)
CONTACT = os.path.join(VAL, 'fig-contact-sheet.jpg')

# ---------------------------------------------------------------- 文件
doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(1.0)
sec.top_margin = sec.bottom_margin = Inches(1.0)
normal = doc.styles['Normal']
normal.font.name = FONT; normal.font.size = Pt(11); normal.font.color.rgb = INK
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

# ---- 封面
para(doc, '', after=40)
para(doc, '第一屆景觀 AI 應用競賽　AI × LANDSCAPE 2026', 9.5, GREY, spacing=20, after=2)
para(doc, '第二組｜設計構想產生', 9.5, GREY, spacing=20, after=26)
para(doc, '植徑　NativePlant AI v3', 16, INK, True, after=6)
para(doc, '附錄 B｜S_attr / S_llm 驗證結果報告', 22, INK, True, after=6)
para(doc, 'APPENDIX B — SIMILARITY SCORING VALIDATION REPORT', 9.5, GREY, spacing=30, after=26)
para(doc, '以視覺語言模型作為外部效標，檢驗規則式屬性評分的候選排序是否成立；'
          '並據以定位目前系統的真實瓶頸。', 11, INK, after=30, line=1.6)

kv_block(doc, [
    ('對應作品', 'zhijing-v3.netlify.app'),
    ('對應章節', '設計過程說明書 §5.2 發散到收斂與方案評估優化'),
    ('驗證版本', 'v.3.1（plants_enriched.json，508 筆）'),
    ('執行日期', '2026 年 8 月 4 日'),
    ('文件性質', '實驗性 prototype 驗證，非最終效能宣稱'),
])

doc.add_page_break()

# ---- 摘要
h1(doc, '摘要　EXECUTIVE SUMMARY')
para(doc, '本輪驗證依設計過程說明書 §5.2 所規劃的方法執行 Phase 1–3，實際完成 40 組「參考植株 × 候選植株」'
          '的雙圖評審，另就 15 組配對各重複評分 3 次以檢驗評審一致性。結論如下：', after=10)
table(doc,
      ['驗證指標', '門檻', '本輪結果', '判定'],
      [['Spearman ρ（S_attr vs S_llm）', 'ρ ≥ 0.6 最低／ρ ≥ 0.7 目標',
        f"組內平均 {P2['summary']['rho_mean']}；全樣本 pooled {pooled}", '未通過'],
       ['Overlap@5', '參考指標（無門檻）', f"平均 {P2['summary']['overlap5_mean']} / 5", '參考'],
       ['S_llm 一致性 SD', 'SD ≤ 15', f"平均 {round(st.mean(sds),2)}（最大 {max(sds)}）", '通過']],
      widths=[2.3, 1.6, 2.0, 0.65])
caption(doc, '表 1　驗證指標總覽。ρ 未達門檻，但 S_llm 自身的重複評分一致性極高。')

para(doc, '關鍵判讀：ρ 未達門檻，不應直接讀作「S_attr 評分失效」。本輪同時發現兩項在相關性之前必須先修復的'
          '前置缺陷——其一為程式層的類別詞彙不一致（§5.1），其二為影像語料的取景與正確性問題（§5.2）。'
          '在效標輸入本身不可靠的情況下，ρ 不具備推翻或支持 S_attr 權重設計的證據力。'
          '因此本輪主動中止 Phase 4（權重校準）與 Phase 5（Role 維度 A/B），避免在受污染的效標上做過適校準。', after=10)

# ---- 一
h1(doc, '一、驗證目的與範圍　SCOPE')
para(doc, '植徑以規則式的 S_attr 屬性加權評分為候選植物排序。此分數透明、可回溯，但其排序是否貼近設計者'
          '實際的視覺判斷，無法由分數本身自證。本驗證引入獨立的視覺語言模型（VLM-as-Judge）作為外部效標 S_llm，'
          '以等級相關係數檢驗兩者的一致性。', after=8)
h2(doc, '本輪涵蓋範圍')
bullet(doc, 'Phase 1：建立可重跑的驗證腳本與 S_attr-only 基準排名')
bullet(doc, 'Phase 2：S_attr vs S_llm 相關性驗證（5 組輸入植株 × 各 8 筆候選 ＝ 40 組配對）')
bullet(doc, 'Phase 3：S_llm 一致性檢驗（5 組輸入 × 各前 3 名候選 × 重複 3 次 ＝ 45 次評分）')
h2(doc, '本輪未執行範圍')
bullet(doc, 'Phase 4 權重校準（NNLS）與 Phase 5 Role 維度 A/B：因效標可靠度不足而主動中止，理由見 §6')
bullet(doc, 'Phase 6 人類專家比對：本期資源未及，仍列為後續優先項')

# ---- 二
h1(doc, '二、方法　METHOD')
h2(doc, '2.1　S_attr 評分結構（總分 100）')
para(doc, '驗證腳本內的評分邏輯與前端 src/matching.js 的 scorePair() 逐行一致，確保驗證對象即為線上實際運行的評分器。', after=8)
table(doc, ['分項', '權重', '計算方式'],
      [['類別 Category', '50（硬性門檻）', '輸入植株與資料庫植株類別須完全相同，否則排除'],
       ['葉形 Foliage', '30', '27 項葉形關鍵字詞庫 → Jaccard 相似度'],
       ['高度 Height', '10', '高度區間與資料庫 HEIGHT 欄位的數值重疊比例'],
       ['展幅 Spread', '5', '展幅區間與資料庫 SPREAD 欄位的數值重疊比例'],
       ['觀賞 Ornament', '5', '26 項觀賞特徵關鍵字（花期＋結構期）→ Jaccard 相似度']],
      widths=[1.35, 1.35, 3.85])
caption(doc, '表 2　S_attr 分項權重。色彩刻意排除於本分數之外，由 Step 02 色彩篩選器獨立處理。')

h2(doc, '2.2　S_llm 評審設計')
para(doc, '以 Claude Haiku 4.5 Vision 為評審模型，對每組配對同時輸入兩張植物照片（圖 A＝參考植株、圖 B＝候選替代植株），'
          '要求其只依「視覺形態相似度」給出 0–100 分並附 40 字以內理由，輸出限定為 JSON。'
          '評分提示詞明確要求評審不得以花色或葉色作為主要依據，以與 S_attr 排除色彩的設計決策對齊。', after=8)
table(doc, ['項目', '設定'],
      [['評審模型', 'claude-haiku-4-5-20251001（Anthropic Messages API）'],
       ['溫度 temperature', '1（保留取樣變異，使 Phase 3 的一致性 SD 具意義）'],
       ['影像處理', '本地照片統一縮至長邊 768 px、JPEG 品質 82'],
       ['評分尺度', '90–100 可直接互換／75–89 高度相似／60–74 中度／40–59 低度／0–39 不相似'],
       ['輸出解析', '正則擷取 JSON；解析失敗自動重試至多 3 次']],
      widths=[1.5, 5.0])
caption(doc, '表 3　S_llm 評審參數。')

h2(doc, '2.3　樣本設計')
para(doc, '採 leave-one-out 設計：自資料庫抽取一筆作為「輸入植株」，以其屬性模擬 Step 02 影像辨識的輸出，'
          '再於資料庫其餘物種中排序。由於 S_llm 需要雙方皆有照片，候選池限縮為本地備有照片的 25 種，'
          '輸入植株則以類別分層、固定亂數種子（seed = 20260804）抽樣，確保可重現。', after=8)
table(doc, ['項目', '數值'],
      [['資料庫總數', '508 種（plants_enriched.json）'],
       ['具本地照片、可進入 S_llm 評審者', '25 種（4.9%）'],
       ['Phase 2 輸入植株', f"{P2['summary']['n_inputs']} 組｜" + '、'.join(r['input']['id'] for r in P2['runs'])],
       ['Phase 2 配對數', f"{P2['summary']['n_pairs']} 組（每組輸入取 S_attr 前 8 名候選）"],
       ['Phase 3 配對數', f"{P3['summary']['n_pairs']} 組 × 重複 3 次 ＝ 45 次評分"],
       ['累計 API 評審次數', '79 次｜實際費用約 US$0.25']],
      widths=[2.35, 4.15])
caption(doc, '表 4　樣本規模。Phase 3 因每組僅取前 3 名，候選池門檻放寬，抽樣結果與 Phase 2 不完全相同。')

h2(doc, '2.4　統計指標')
bullet(doc, 'Spearman ρ：組內計算（每組輸入植株的 8 筆候選各自排名），同分採平均秩；另計算全樣本 pooled ρ 供對照')
bullet(doc, 'Overlap@5：S_attr 前 5 名與 S_llm 前 5 名的交集筆數')
bullet(doc, 'S_llm 一致性 SD：同一配對重複 3 次評分的樣本標準差')

# ---- 三
h1(doc, '三、結果　RESULTS')
h2(doc, '3.1　Phase 2 主結果')
rows = []
for r in P2['runs']:
    rows.append([r['input']['id'], r['input']['category'], r['n_pool'], len(r['candidates']),
                 r['rho'], f"{r['overlap5']} / 5"])
rows.append(['平均', '—', '—', P2['summary']['n_pairs'], P2['summary']['rho_mean'],
             f"{P2['summary']['overlap5_mean']} / 5"])
table(doc, ['輸入植株', '類別', '候選池', '送審', 'Spearman ρ', 'Overlap@5'], rows,
      widths=[1.5, 0.8, 0.8, 0.7, 1.35, 1.35])
caption(doc, '表 5　各組輸入植株的排序一致性。5 組中有 3 組 ρ 為負，平均值接近零。')

para(doc, f"全樣本 pooled ρ ＝ {pooled}，與組內平均一致地接近零；顯示在目前的資料條件下，"
          f"S_attr 的排序與 S_llm 的排序之間不存在可辨識的單調關係。", after=10)

doc.add_picture(SCATTER, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption(doc, '圖 1　S_attr 對 S_llm 散布圖（40 組配對）。橫軸為規則式評分、縱軸為 VLM 評審分數。'
             '點群沿縱軸大幅壓縮於 60 分以下，且未呈現正向趨勢。')

h2(doc, '3.2　S_llm 分數分布')
table(doc, ['統計量', '數值'],
      [['樣本數', f"{len(S_llm)} 組配對"],
       ['最小值 / 最大值', f"{min(S_llm)} / {max(S_llm)}"],
       ['平均數 / 中位數', f"{round(st.mean(S_llm),1)} / {st.median(S_llm)}"],
       ['標準差', f"{round(st.pstdev(S_llm),1)}"],
       ['≥ 60 分（中度相似以上）', f"{sum(1 for v in S_llm if v>=60)} / {len(S_llm)}（{round(100*sum(1 for v in S_llm if v>=60)/len(S_llm))}%）"],
       ['≤ 45 分（低度相似以下）', f"{sum(1 for v in S_llm if v<=45)} / {len(S_llm)}（{round(100*sum(1 for v in S_llm if v<=45)/len(S_llm))}%）"]],
      widths=[2.4, 4.1])
caption(doc, '表 6　S_llm 分數分布。逾八成配對落在 45 分以下，呈明顯地板效應；'
             '評審實際只使用了尺度的下半段，可鑑別的層級因而不足。')

h2(doc, '3.3　S_attr 分項與 S_llm 的相關性')
label = {'foliage': '葉形 Foliage（權重 30）', 'height': '高度 Height（權重 10）',
         'spread': '展幅 Spread（權重 5）', 'ornament': '觀賞 Ornament（權重 5）'}
table(doc, ['S_attr 分項', '可計算組數', '組內平均 ρ'],
      [[label[k], f"{v[0]} / 5", f"{v[1]:+.3f}" if v[1] is not None else '—'] for k, v in comp_rho.items()],
      widths=[2.6, 1.5, 2.4])
caption(doc, '表 7　各分項與 S_llm 的組內相關。四個分項全部接近零；觀賞分項在 5 組中有 4 組所有候選同分，'
             '完全不具鑑別力，僅 1 組可計算。')

h2(doc, '3.4　Phase 3　S_llm 一致性')
rows3 = []
for r in P3['runs']:
    for c in r['candidates']:
        rows3.append([r['input']['id'], c['name'], c['s_attr'],
                      ' / '.join(str(v) for v in c['s_llm_runs']), c['s_llm'], c['s_llm_sd']])
table(doc, ['輸入植株', '候選植株', 'S_attr', '三次評分', '平均', 'SD'], rows3,
      widths=[1.25, 1.25, 0.75, 1.25, 0.7, 0.6], size=9)
caption(doc, f"表 8　重複評分一致性。平均 SD ＝ {round(st.mean(sds),2)}，遠低於門檻 15，"
             f"其中 {sum(1 for v in sds if v==0)} 組配對三次評分完全相同。")

para(doc, 'S_llm 通過一致性檢驗，代表評審本身穩定、不是隨機亂給分。但須留意：這份高一致性有相當部分來自'
          '分數集中於 28–35 的窄帶——穩定地給出相近的低分，同樣會產生極低的 SD。'
          '因此本項通過只能說明「評審可重複」，不能推論「評審有鑑別力」。', after=10)

# ---- 四
h1(doc, '四、發現與診斷　FINDINGS')

h2(doc, '4.1　阻斷性缺陷：類別詞彙不一致（已於驗證腳本修正）')
para(doc, 'scorePair() 以 input.category === db.category 作為硬性門檻。但兩端的詞彙表並不相同：', after=6)
table(doc, ['來源', '類別詞彙'],
      [['analyze-image.js 影像辨識輸出', '灌木／草本／地被'],
       ['plants_enriched.json 資料庫欄位', '灌木（133）／喬木（133）／多年生草本（211）／一年生草本（31）']],
      widths=[2.6, 3.9])
caption(doc, '表 9　兩端類別詞彙對照。')
para(doc, '字串完全相等的比對下，只有「灌木」一類能成功配對；辨識為「草本」或「地被」的植株，'
          '在 508 筆資料中永遠得不到任何候選，池子為空。此缺陷影響全系統約半數的辨識結果，'
          '且在 Demo 中會表現為「該類別無符合植物」，不易被察覺為程式錯誤。', after=8)
para(doc, '驗證腳本已加入 normCategory() 做映射：灌木→灌木；草本類依 Role 欄位再細分，'
          'Ground cover→地被，其餘→草本；喬木不參與比對。建議將同一函式回寫 src/matching.js。', after=10)

h2(doc, '4.2　主要瓶頸：影像語料的取景與正確性')
para(doc, 'S_llm 的輸入是照片，因此照片品質直接決定效標品質。本輪對 25 張實際送審影像逐張稽核，結果如下：', after=8)
table(doc, ['取景類型', '張數', '占比', '對形態比對的可用性'],
      [['全株／群落尺度', '4', '16%', '可判讀株形、質地與量體'],
       ['枝條／植株局部', '4', '16%', '部分可判讀'],
       ['花／果／單葉特寫', '15', '60%', '無法判讀株形與展幅'],
       ['主體非目標植物', '2', '8%', '不可用']],
      widths=[1.65, 0.7, 0.7, 3.45])
caption(doc, '表 10　25 張送審影像的取景稽核。')
bullet(doc, '檔名標示 _habit（全株）者共 17 張，其中僅 4 張為真正的全株照，名實不符率 76%')
bullet(doc, '青葙_habit.jpg 的主體為一隻鳥（Lilac-breasted Roller），完全非目標植物')
bullet(doc, '綬草_habit.jpg 為古典植物版畫，非攝影照片')
bullet(doc, '燈稱花_habit.jpg 畫面中含園區中文解說牌；絡石、高士拂澤蘭的照片中昆蟲占顯著面積；'
            '田代氏石斑木_habit.jpg 為住宅街景，含房屋、柵欄與垃圾桶')
para(doc, '此問題在評審理由中留下了直接證據。以青葙為輸入或候選的配對，評審一律回覆「圖 B 為鳥類，非植物，'
          '無法進行植物替代評估」並給出最低分；以台灣金絲桃為輸入的 8 組配對中，評審 7 次將其誤判為「蕨類」。'
          '40 組配對中，共 11 組（27.5%）的評審理由顯示其對影像主體的判讀與資料庫標定不符。', after=10)
if os.path.exists(CONTACT):
    doc.add_picture(CONTACT, width=Inches(5.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, '圖 2　25 張送審影像接觸表（編號 1–25 依物種中文名排序）。第 22 張為鳥類、第 17 張為植物版畫、'
                 '第 12 張含解說牌、第 13 張為住宅街景；多數影像為花果特寫而非株形。')

h2(doc, '4.3　S_attr 側的鑑別力不足')
para(doc, '影像之外，S_attr 本身在小候選池上的分辨能力也偏弱。40 組候選中，★3 占 26 組（65%），'
          '★5 僅 1 組；分數集中於 51.3–90 的窄帶（標準差 8.4），其中 50 分為類別門檻的固定基底。'
          '換言之，實際產生排序的僅有 50 分的軟性區間，而觀賞分項在多數組別完全同分。'
          '在此情形下，即使效標無誤，ρ 的統計不確定性也偏高。', after=10)

# ---- 五
h1(doc, '五、結論　CONCLUSION')
para(doc, '本輪驗證的判定為：Spearman ρ 未達 0.6 的最低可接受門檻；S_llm 一致性 SD 通過門檻。', after=8)
para(doc, '但依 §4 的診斷，此結果不足以推論 S_attr 的權重設計有誤。理由是效標鏈上游存在兩處已確認的污染：'
          '類別詞彙不一致使可比對範圍失真，而 60% 為花果特寫、8% 主體錯誤的影像語料，'
          '使 S_llm 有相當比例是在比較「錯誤的東西」。'
          '在此條件下對權重做 NNLS 校準，只會把權重擬合到影像缺陷上，'
          '這正是 Phase 4 保留樣本設計原本要防範的過適風險。故本輪主動中止 Phase 4 與 Phase 5。', after=8)
para(doc, '就 prototype 的目的而言，本輪的價值不在於取得一個漂亮的 ρ，而在於驗證流程本身已可運作、'
          '且成功地把系統的真實瓶頸從「權重是否正確」重新定位到「資料與影像是否可用」。'
          '這是一個比相關係數更有行動意義的結論。', after=10)

h1(doc, '六、後續行動　NEXT ACTIONS')
table(doc, ['優先序', '行動', '預期效果'],
      [['P0', '將 normCategory() 回寫 src/matching.js，並補上類別對照的單元測試',
        '修復草本／地被兩類永遠無候選的阻斷性缺陷'],
       ['P0', '重建影像語料：每種至少一張真正的全株照，剔除主體錯誤與含人工物、動物的影像，'
              '並建立 _habit／_close／_leaf 三槽的驗收規則',
        '解除 S_llm 效標的主要污染來源'],
       ['P1', '語料修復後重跑 Phase 2（建議擴至 10–15 組輸入），重新判定 ρ',
        '取得可信的相關性基準'],
       ['P1', '檢討觀賞 Ornament 分項的關鍵字詞庫或權重',
        '改善多數組別同分、無鑑別力的問題'],
       ['P2', '待 ρ 達 0.6 以上再執行 Phase 4 權重校準與 Phase 5 Role 維度 A/B',
        '避免在受污染效標上過適'],
       ['P2', 'Phase 6：以 2–3 位景觀設計師對 12–15 組配對評分，驗證 S_llm 作為效標的可信度',
        '補強效標本身的效度論證']],
      widths=[0.65, 3.15, 2.7])
caption(doc, '表 11　後續行動優先序。P0 為重跑驗證前必須完成的前置修復。')

h1(doc, '七、可重現性　REPRODUCIBILITY')
para(doc, '本報告的所有數值皆由下列指令產生，固定亂數種子後可完整重現：', after=8)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
set_font(p.add_run('node scripts/validate-similarity.mjs --inputs 5 --topk 8 --repeat 1 \\\n'
                   '     --model claude-haiku-4-5-20251001 --seed 20260804'), 9.5, INK)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
set_font(p.add_run('node scripts/validate-similarity.mjs --inputs 5 --topk 3 --repeat 3 \\\n'
                   '     --model claude-haiku-4-5-20251001 --seed 20260804'), 9.5, INK)
table(doc, ['檔案', '內容'],
      [['scripts/validate-similarity.mjs', '驗證腳本（S_attr 評分、抽樣、S_llm 呼叫、統計）'],
       ['scripts/build-validation-report.py', '本報告的產生腳本'],
       ['validation/phase2-results.json', 'Phase 2 完整結果（含每筆評審理由）'],
       ['validation/phase3-results.json', 'Phase 3 重複評分結果'],
       ['validation/phase2-pairs.csv', '40 組配對的分項明細（可直接開啟檢視）'],
       ['validation/.judge-cache.json', '評審結果落地快取，重跑不重複計費']],
      widths=[2.5, 4.0])
caption(doc, '表 12　產出檔案清單。')

h1(doc, '八、AI 工具揭露　AI DISCLOSURE')
table(doc, ['角色', '模型 / 版本', '用途'],
      [['視覺相似度評審', 'claude-haiku-4-5-20251001', 'S_llm 效標，對雙圖配對給出 0–100 分'],
       ['影像辨識（系統內）', 'GPT-4o Vision', '本驗證以資料庫屬性模擬其輸出，未實際呼叫'],
       ['規則式評分', '無 AI（確定性程式碼）', 'S_attr，與 src/matching.js 邏輯一致']],
      widths=[1.6, 2.2, 2.7])
para(doc, '完整提示詞、模型版本、溫度設定與每筆評審的原始回覆，均保存於 validation/phase2-results.json 與'
          ' scripts/validate-similarity.mjs，可供查驗。', 9.5, GREY, after=6)

out = os.path.join(VAL, '植徑v3_附錄B_S_attr_S_llm_驗證結果報告.docx')
doc.save(out)
print('saved:', out)
